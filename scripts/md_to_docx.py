#!/usr/bin/env python3
"""Convert a markdown report into a styled .docx.

Usage: python md_to_docx.py <input.md> <output.docx>

Handles headings, tables, code blocks, bullets, numbered items, quotes,
bold/italic/code inline markers and horizontal rules.
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
LINK_RE = re.compile(r"\[([^\]]+)\]\([^)]*\)")

HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
BULLET_RE = re.compile(r"^(\s*)[-*+]\s+(.*)$")
NUMBERED_RE = re.compile(r"^\s*\d+\.\s+(.*)$")
QUOTE_RE = re.compile(r"^>\s?(.*)$")
HR_RE = re.compile(r"^\s*(?:-{3,}|\*{3,}|_{3,})\s*$")

HEADING_COLORS = {
    1: RGBColor(0x1F, 0x3B, 0x5C),
    2: RGBColor(0x1F, 0x4E, 0x79),
    3: RGBColor(0x2E, 0x5E, 0x8C),
    4: RGBColor(0x2E, 0x5E, 0x8C),
    5: RGBColor(0x2E, 0x5E, 0x8C),
    6: RGBColor(0x2E, 0x5E, 0x8C),
}


def add_shading(paragraph, fill="F2F2F2"):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def add_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def inline_runs(paragraph, text, bold=False, size=None):
    text = LINK_RE.sub(r"\1", text)
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if len(part) > 4 and part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif len(part) > 2 and part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10 if size is None else size)
        elif len(part) > 2 and part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        else:
            r = paragraph.add_run(part)
            r.bold = bold


def is_table_line(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|")


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def is_separator_row(cells):
    return bool(cells) and all(re.fullmatch(r":?-{1,}:?", c) for c in cells)


def add_page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def main():
    if len(sys.argv) != 3:
        print("usage: python md_to_docx.py <input.md> <output.docx>")
        sys.exit(2)
    md_path = Path(sys.argv[1])
    out_path = Path(sys.argv[2])

    lines = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    add_page_number_footer(doc)

    i, n = 0, len(lines)
    first_h1 = True
    code_buf = []

    def flush_code():
        nonlocal code_buf
        if not code_buf:
            return
        for cl in code_buf:
            p = doc.add_paragraph()
            r = p.add_run(cl if cl else " ")
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            add_shading(p)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Pt(6)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(8)
        sp.add_run("").font.size = Pt(2)
        code_buf = []

    def spacer():
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after = Pt(6)
        sp.add_run("").font.size = Pt(2)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Code fence
        if stripped.startswith("```"):
            code_buf = []
            i += 1
            while i < n and not lines[i].strip().startswith("```"):
                code_buf.append(lines[i])
                i += 1
            i += 1  # closing fence
            flush_code()
            continue

        # Table
        if is_table_line(line):
            rows = []
            while i < n and is_table_line(lines[i]):
                rows.append(split_row(lines[i]))
                i += 1
            rows = [r for r in rows if not is_separator_row(r)]
            if rows:
                width = max(len(r) for r in rows)
                rows = [r + [""] * (width - len(r)) for r in rows]
                table = doc.add_table(rows=len(rows), cols=width)
                table.style = "Table Grid"
                table.autofit = True
                for ri, row in enumerate(rows):
                    for ci, cell_text in enumerate(row):
                        cell = table.cell(ri, ci)
                        cell.text = ""
                        inline_runs(cell.paragraphs[0], cell_text, bold=(ri == 0))
                        if ri == 0:
                            add_cell_shading(cell, "DCE6F1")
                spacer()
            continue

        if not stripped:
            i += 1
            continue

        m = HEADING_RE.match(line)
        if m:
            level = len(m.group(1))
            text = LINK_RE.sub(r"\1", m.group(2)).strip()
            if level == 1 and first_h1:
                h = doc.add_heading("", level=0)
                inline_runs(h, text)
                for r in h.runs:
                    r.font.color.rgb = HEADING_COLORS[1]
                first_h1 = False
            else:
                h = doc.add_heading("", level=level)
                inline_runs(h, text)
                for r in h.runs:
                    r.font.color.rgb = HEADING_COLORS.get(level, RGBColor(0x2E, 0x5E, 0x8C))
            i += 1
            continue

        if HR_RE.match(line):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_shading(p, "FFFFFF")
            i += 1
            continue

        m = BULLET_RE.match(line)
        if m:
            indent = len(m.group(1)) // 2
            style = "List Bullet" if indent < 2 else "List Bullet 2"
            p = doc.add_paragraph(style=style)
            inline_runs(p, m.group(2))
            i += 1
            continue

        m = NUMBERED_RE.match(line)
        if m:
            p = doc.add_paragraph()
            inline_runs(p, line.strip())
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        m = QUOTE_RE.match(line)
        if m:
            p = doc.add_paragraph()
            inline_runs(p, m.group(1))
            for r in p.runs:
                r.italic = True
                r.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
            p.paragraph_format.left_indent = Pt(18)
            i += 1
            continue

        p = doc.add_paragraph()
        inline_runs(p, line)
        i += 1

    doc.save(out_path)
    print(f"saved {out_path} ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
